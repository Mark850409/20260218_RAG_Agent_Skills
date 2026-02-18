"""
多格式文件載入器
支援：Markdown (.md)、PDF (.pdf)、Word (.docx)、CSV (.csv)、Excel (.xlsx)
"""

import os
import csv
from pathlib import Path
from typing import List, Dict, Any


def load_document(file_path: str) -> List[Dict[str, Any]]:
    """
    載入文件並返回文字片段列表
    每個片段包含：text（文字內容）、metadata（來源資訊）
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    loaders = {
        ".md": _load_markdown,
        ".txt": _load_text,
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".csv": _load_csv,
        ".xlsx": _load_excel,
        ".xls": _load_excel,
    }

    if ext not in loaders:
        raise ValueError(f"不支援的文件格式：{ext}，支援格式：{', '.join(loaders.keys())}")

    return loaders[ext](file_path)


def _build_chunk(text: str, source: str, page: int = 0, extra: Dict = None) -> Dict[str, Any]:
    """建立標準化的文字片段"""
    metadata = {"source": source, "filename": Path(source).name, "page": page}
    if extra:
        metadata.update(extra)
    return {"text": text.strip(), "metadata": metadata}


def _load_markdown(file_path: str) -> List[Dict[str, Any]]:
    """載入 Markdown 文件，依標題分割"""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()

    chunks = []
    current_section = []
    current_title = "開頭"

    for line in content.split("\n"):
        if line.startswith("#"):
            # 儲存前一個段落
            if current_section:
                text = "\n".join(current_section).strip()
                if text:
                    chunks.append(_build_chunk(text, file_path, extra={"section": current_title}))
            current_title = line.lstrip("#").strip()
            current_section = [line]
        else:
            current_section.append(line)

    # 儲存最後一個段落
    if current_section:
        text = "\n".join(current_section).strip()
        if text:
            chunks.append(_build_chunk(text, file_path, extra={"section": current_title}))

    # 若無標題結構，整份文件作為一個 chunk
    if not chunks:
        chunks.append(_build_chunk(content, file_path))

    return chunks


def _load_text(file_path: str) -> List[Dict[str, Any]]:
    """載入純文字文件"""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    return [_build_chunk(content, file_path)]


def _load_pdf(file_path: str) -> List[Dict[str, Any]]:
    """載入 PDF 文件，每頁為一個 chunk"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("請安裝 pypdf：pip install pypdf")

    reader = PdfReader(file_path)
    chunks = []

    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            chunks.append(_build_chunk(text, file_path, page=page_num + 1))

    if not chunks:
        chunks.append(_build_chunk("（PDF 無法提取文字內容）", file_path))

    return chunks


def _load_docx(file_path: str) -> List[Dict[str, Any]]:
    """載入 Word 文件，依段落分組"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("請安裝 python-docx：pip install python-docx")

    doc = Document(file_path)
    chunks = []
    current_section = []
    current_heading = "開頭"

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            if current_section:
                text = "\n".join(current_section).strip()
                if text:
                    chunks.append(_build_chunk(text, file_path, extra={"section": current_heading}))
            current_heading = para.text.strip() or current_heading
            current_section = []
        else:
            if para.text.strip():
                current_section.append(para.text)

    if current_section:
        text = "\n".join(current_section).strip()
        if text:
            chunks.append(_build_chunk(text, file_path, extra={"section": current_heading}))

    if not chunks:
        all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        chunks.append(_build_chunk(all_text, file_path))

    return chunks


def _load_csv(file_path: str) -> List[Dict[str, Any]]:
    """載入 CSV 文件，每 N 行為一個 chunk"""
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("請安裝 pandas：pip install pandas")

    try:
        df = pd.read_csv(file_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding="big5", errors="ignore")

    chunks = []
    columns = list(df.columns)
    header = "欄位：" + "、".join(str(c) for c in columns)

    # 每 20 行為一個 chunk
    batch_size = 20
    for i in range(0, len(df), batch_size):
        batch = df.iloc[i : i + batch_size]
        rows_text = []
        for _, row in batch.iterrows():
            row_text = "；".join(f"{col}={val}" for col, val in zip(columns, row.values) if str(val) != "nan")
            rows_text.append(row_text)

        text = f"{header}\n" + "\n".join(rows_text)
        chunks.append(_build_chunk(text, file_path, extra={"rows": f"{i+1}-{min(i+batch_size, len(df))}"}))

    return chunks if chunks else [_build_chunk("（CSV 檔案為空）", file_path)]


def _load_excel(file_path: str) -> List[Dict[str, Any]]:
    """載入 Excel 文件，每個工作表分別處理"""
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("請安裝 pandas 與 openpyxl：pip install pandas openpyxl")

    xl = pd.ExcelFile(file_path)
    chunks = []

    for sheet_name in xl.sheet_names:
        try:
            df = xl.parse(sheet_name)
        except Exception:
            continue

        if df.empty:
            continue

        columns = list(df.columns)
        header = f"工作表：{sheet_name}，欄位：" + "、".join(str(c) for c in columns)

        # 每 20 行為一個 chunk
        batch_size = 20
        for i in range(0, len(df), batch_size):
            batch = df.iloc[i : i + batch_size]
            rows_text = []
            for _, row in batch.iterrows():
                row_text = "；".join(
                    f"{col}={val}" for col, val in zip(columns, row.values) if str(val) not in ("nan", "NaT", "")
                )
                if row_text:
                    rows_text.append(row_text)

            if rows_text:
                text = f"{header}\n" + "\n".join(rows_text)
                chunks.append(
                    _build_chunk(text, file_path, extra={"sheet": sheet_name, "rows": f"{i+1}-{min(i+batch_size, len(df))}"})
                )

    return chunks if chunks else [_build_chunk("（Excel 檔案為空）", file_path)]


def get_supported_extensions() -> List[str]:
    """返回支援的文件格式列表"""
    return [".md", ".txt", ".pdf", ".docx", ".csv", ".xlsx", ".xls"]
